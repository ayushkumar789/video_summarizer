from docx import Document

def generate_word(transcript, summary, keyframes, output_path):
    doc = Document()
    doc.add_heading("Video Transcription", level=1)
    doc.add_paragraph(transcript)

    doc.add_heading("Summary", level=1)
    doc.add_paragraph(summary)

    doc.add_heading("Keyframes", level=1)
    for frame in keyframes:
        doc.add_paragraph(f"Keyframe: {frame}")

    doc.save(output_path)
